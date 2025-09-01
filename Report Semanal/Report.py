import os
import pandas as pd
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def selecionar_arquivo_excel():
    """Seleciona arquivo Excel da pasta atual."""
    arquivos = []
    for arquivo in os.listdir():
        if arquivo.endswith(('.xls', '.xlsx')):
            arquivos.append(arquivo)
    
    if not arquivos:
        raise FileNotFoundError("Nenhum arquivo Excel encontrado na pasta atual.")
    elif len(arquivos) == 1:
        print(f"Arquivo selecionado automaticamente: {arquivos[0]}")
        return arquivos[0]
    else:
        print("Selecione um arquivo:")
        for idx, arquivo in enumerate(arquivos, 1):
            print(f"{idx}. {arquivo}")
        
        while True:
            try:
                escolha = int(input("Digite o número do arquivo desejado: "))
                if 1 <= escolha <= len(arquivos):
                    return arquivos[escolha - 1]
                else:
                    print(f"Por favor, digite um número entre 1 e {len(arquivos)}")
            except ValueError:
                print("Por favor, digite apenas números")
            except KeyboardInterrupt:
                print("\nOperação cancelada pelo usuário.")
                exit(1)


# Dicionário para tradução de meses
meses_portugues = {
    'janeiro': 'January', 'fevereiro': 'February', 'março': 'March', 'abril': 'April',
    'maio': 'May', 'junho': 'June', 'julho': 'July', 'agosto': 'August',
    'setembro': 'September', 'outubro': 'October', 'novembro': 'November', 'dezembro': 'December'
}


def traduzir_meses(texto):
    """Traduz nomes de meses do português para inglês."""
    if pd.isna(texto):
        return texto
    
    texto_lower = str(texto).lower()
    for pt, en in meses_portugues.items():
        texto_lower = texto_lower.replace(pt, en)
    return texto_lower


def formatar_data(coluna):
    """Formata coluna de datas para o padrão brasileiro."""
    try:
        datas_traduzidas = coluna.apply(traduzir_meses)
        datas_convertidas = pd.to_datetime(
            datas_traduzidas, 
            format='%d %B %Y %H:%M', 
            errors='coerce'
        )
        return datas_convertidas.dt.strftime('%d/%m/%y')
    except Exception as e:
        print(f"Aviso: Erro ao formatar datas: {e}")
        return coluna


def processar_porcentagem(valor):
    """Processa valor de porcentagem removendo símbolos e convertendo."""
    if pd.isna(valor):
        return 0
    
    valor_str = str(valor).replace('%', '').replace(',', '.')
    try:
        return float(valor_str) / 100 if '%' in str(valor) else float(valor_str)
    except (ValueError, TypeError):
        return 0


def calcular_dias_diferenca(data1, data2):
    """Calcula diferença em dias entre duas datas."""
    if pd.isna(data1) or pd.isna(data2):
        return 0
    try:
        return (data1 - data2).days
    except:
        return 0


def buscar_hierarquia(df, linha_index):
    """Busca a hierarquia (bisavô, avô, pai) de uma linha específica."""
    pai = avo = bisavo = ''
    
    for i in range(linha_index - 1, -1, -1):
        if i < 0 or i >= len(df):
            continue
            
        try:
            nivel = df.at[i, 'Nível_da_estrutura_de_tópicos']
            nome = str(df.at[i, 'Nome'])
            
            if nivel == 3 and not pai:
                pai = nome
            elif nivel == 2 and not avo:
                avo = nome
            elif nivel == 1 and not bisavo:
                bisavo = nome
            
            if pai and avo and bisavo:
                break
        except (KeyError, IndexError):
            continue
    
    return bisavo, avo, pai


def filtrar_tarefas_por_recurso(df, termo_busca):
    """Filtra tarefas baseado no recurso especificado."""
    try:
        filtro_recursos = df['Nomes_dos_Recursos'].astype(str).str.contains(termo_busca, case=False, na=False)
        filtro_percentual = (df['Porcentagem_Concluída'] > 0) & (df['Porcentagem_Concluída'] < 1)
        return df[filtro_recursos & filtro_percentual]
    except KeyError:
        print(f"Aviso: Coluna necessária não encontrada para filtrar {termo_busca}")
        return pd.DataFrame()


def adicionar_secao_documento(doc, titulo, tarefas_df, df_principal, hoje, tipo_secao):
    """Adiciona uma seção ao documento Word com as tarefas especificadas."""
    doc.add_paragraph(f"\n{titulo}")
    
    if tarefas_df.empty:
        doc.add_paragraph("- Não existem tarefas que cumpram os critérios desta seção")
        return
    
    # Agrupar tarefas por hierarquia
    grupos = {}
    for idx, row in tarefas_df.iterrows():
        bisavo, avo, pai = buscar_hierarquia(df_principal, idx)
        chave = bisavo if bisavo else "Sem categoria"
        
        if tipo_secao == "emissoes":
            linha = f"{avo} - {pai} - {row['Nome']}: Programado para {row.get('Término', 'N/A')}"
        else:  # analise
            dias_analise = "?"
            if pd.notna(row.get('Início_DT')):
                try:
                    dias_analise = (hoje - row['Início_DT']).days
                except:
                    dias_analise = "?"
            linha = f"{avo} - {pai} - {row['Nome']}: Com o cliente desde {row.get('Início', 'N/A')} ({dias_analise} dias)"
        
        if chave not in grupos:
            grupos[chave] = []
        grupos[chave].append(linha)
    
    # Adicionar grupos ao documento
    for bisavo, tarefas in grupos.items():
        if bisavo:
            doc.add_paragraph(f"\n{bisavo}:")
        for tarefa in tarefas:
            p = doc.add_paragraph(tarefa)
            p.style = 'List Bullet'


def validar_colunas_necessarias(df):
    """Valida se as colunas necessárias existem no DataFrame."""
    colunas_obrigatorias = [
        'Nível_da_estrutura_de_tópicos', 'Nome', 'Nomes_dos_Recursos',
        'Porcentagem_Concluída', 'Porcentagem_Previsto'
    ]
    
    colunas_faltantes = []
    for coluna in colunas_obrigatorias:
        if coluna not in df.columns:
            colunas_faltantes.append(coluna)
    
    if colunas_faltantes:
        print(f"Aviso: Colunas não encontradas: {', '.join(colunas_faltantes)}")
        # Criar colunas com valores padrão
        for coluna in colunas_faltantes:
            if 'Porcentagem' in coluna:
                df[coluna] = 0
            else:
                df[coluna] = ''
    
    return df


def gerar_relatorio(nome_projeto):
    """Função principal para gerar o relatório semanal."""
    try:
        print("Iniciando geração do relatório...")
        
        # Selecionar e carregar arquivo
        arquivo = selecionar_arquivo_excel()
        print(f"Carregando dados de: {arquivo}")
        df = pd.read_excel(arquivo)
        
        if df.empty:
            raise ValueError("O arquivo Excel está vazio")
        
        # Validar colunas necessárias
        df = validar_colunas_necessarias(df)
        
        # Processar colunas de data
        col_datas = ['Início', 'Término', 'Início_da_Linha_de_Base', 'Término_da_linha_de_base']
        for col in col_datas:
            if col in df.columns:
                df[col] = formatar_data(df[col])
                # Criar versões datetime das colunas
                df[col + '_DT'] = pd.to_datetime(df[col], format='%d/%m/%y', errors='coerce')
        
        # Processar porcentagens
        if 'Porcentagem_Previsto' in df.columns:
            df['Porcentagem_Previsto'] = df['Porcentagem_Previsto'].apply(processar_porcentagem)
        
        # Obter data atual
        hoje = datetime.now()
        hoje_fmt = hoje.strftime('%d/%m/%y')
        
        # Encontrar linha de nível 0 (projeto principal)
        nivel0_linhas = df[df['Nível_da_estrutura_de_tópicos'] == 0]
        if nivel0_linhas.empty:
            print("Aviso: Nenhuma linha de nível 0 encontrada. Usando primeira linha.")
            nivel0 = df.iloc[0]
        else:
            nivel0 = nivel0_linhas.iloc[0]
        
        # Calcular métricas do projeto
        AA = nivel0.get('Término', 'N/A')
        BB = calcular_dias_diferenca(nivel0.get('Término_DT'), nivel0.get('Término_da_linha_de_base_DT'))
        CC = nivel0.get('Término_da_linha_de_base', 'N/A')
        DD = calcular_dias_diferenca(nivel0.get('Término_da_linha_de_base_DT'), nivel0.get('Início_da_Linha_de_Base_DT'))
        EE = calcular_dias_diferenca(nivel0.get('Término_DT'), nivel0.get('Início_DT'))
        
        # Calcular aderência
        FF = 0
        if nivel0.get('Porcentagem_Previsto', 0) > 0:
            FF = nivel0.get('Porcentagem_Concluída', 0) / nivel0['Porcentagem_Previsto']
        FF_fmt = f"{FF:.0%}"
        
        # Filtrar tarefas
        filtro_horizontes = filtrar_tarefas_por_recurso(df, "Horizontes")
        filtro_cliente = filtrar_tarefas_por_recurso(df, "Cliente")
        
        # Criar documento Word
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(10)
        
        # Título do relatório
        p = doc.add_paragraph()
        run = p.add_run(f"REPORT SEMANAL {nome_projeto.upper()} - {hoje_fmt}")
        run.underline = True
        run.bold = True
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        
        # Seção Resumo
        doc.add_paragraph("📌 RESUMO:")
        
        resumo_textos = [
            f"O projeto, com tendência de término para {AA}, está {abs(BB)} dias corridos {'atrasado' if BB > 0 else 'adiantado'} em relação à Linha de Base aprovada pelo cliente, que previa término em {CC}.",
            f"Com duração inicial de {DD} dias corridos, o projeto possui atualmente duração estimada de {EE} dias corridos.",
            f"O grau de aderência do projeto ao planejamento é de {FF_fmt}."
        ]
        
        for texto in resumo_textos:
            p = doc.add_paragraph(texto)
            p.style = 'List Bullet'
        
        # Adicionar seções
        adicionar_secao_documento(
            doc, "📅 PRÓXIMAS EMISSÕES DE PROJETO:", 
            filtro_horizontes, df, hoje, "emissoes"
        )
        
        adicionar_secao_documento(
            doc, "🔎 ARQUIVOS EM ANÁLISE:", 
            filtro_cliente, df, hoje, "analise"
        )
        
        # Salvar arquivo
        nome_arquivo = f"Relatorio_Semanal_{nome_projeto.replace(' ', '_')}_{hoje_fmt.replace('/', '-')}.docx"
        doc.save(nome_arquivo)
        print(f"\nRelatório salvo como: {nome_arquivo}")
        
    except FileNotFoundError as e:
        print(f"Erro: {e}")
    except Exception as e:
        print(f"Erro inesperado: {e}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    try:
        projeto = input("Digite o nome do projeto: ").strip()
        if not projeto:
            print("Nome do projeto é obrigatório.")
            exit(1)
        gerar_relatorio(projeto)
    except KeyboardInterrupt:
        print("\nOperação cancelada pelo usuário.")
    except Exception as e:
        print(f"Erro na execução: {e}")